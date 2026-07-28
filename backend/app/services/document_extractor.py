import docx
import fitz  # PyMuPDF
from fastapi import HTTPException, status


def get_page_count(file_content: bytes, mime_type: str) -> int:
    try:
        if mime_type == "application/pdf":
            doc = fitz.open(stream=file_content, filetype="pdf")
            count = doc.page_count
            doc.close()
            return count

        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            # Word documents do not have a reliable page count without rendering
            # Estimate based on paragraph count - conservative approximation
            import io

            document = docx.Document(io.BytesIO(file_content))
            paragraph_count = len(document.paragraphs)
            return max(1, paragraph_count // 20)

        elif mime_type == "text/plain":
            # Estimate pages based on character count - ~3000 chars per page
            return max(1, len(file_content) // 3000)

        return 1

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not read document: {str(e)}",
        )


def extract_text(file_content: bytes, mime_type: str) -> str:
    try:
        if mime_type == "application/pdf":
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text

        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            import io

            document = docx.Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)

        elif mime_type == "text/plain":
            return file_content.decode("utf-8", errors="ignore")

        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type for text extraction.",
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to extract text from document: {str(e)}",
        )


def extract_pages(file_content: bytes, mime_type: str) -> list[dict]:
    """
    Returns a list of {"page_number": int, "text": str} for each page.
    For PDFs, page_number is the real page. For formats without real pages
    (Word, plain text), everything is returned as a single page_number=None entry
    so the caller can still process the text while knowing no page data exists.
    """
    try:
        if mime_type == "application/pdf":
            doc = fitz.open(stream=file_content, filetype="pdf")
            pages = []
            for index, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    pages.append(
                        {
                            "page_number": index + 1,  # 1-based, human-readable
                            "text": page_text,
                        }
                    )
            doc.close()
            return pages
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            import io

            document = docx.Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            # Word has no reliable page numbers - return as a single pageless entry
            return [{"page_number": None, "text": full_text}]
        elif mime_type == "text/plain":
            text = file_content.decode("utf-8", errors="ignore")
            return [{"page_number": None, "text": text}]
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type for text extraction.",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to extract text from document: {str(e)}",
        )
